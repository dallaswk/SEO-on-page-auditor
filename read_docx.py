import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document(r"d:\Proyectos\Personal\SeoOnPage\Checklist_Auditoria_SEO_Tecnica.docx")

with open(r"d:\Proyectos\Personal\SeoOnPage\docx_structure.txt", "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPHS ===\n")
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else "NoStyle"
        f.write(f"[{i}] STYLE={style_name} TEXT={p.text[:200]}\n")

    f.write("\n=== TABLES ===\n")
    for ti, table in enumerate(doc.tables):
        f.write(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols\n")
        for ri, row in enumerate(table.rows):
            cells = [c.text[:80].replace('\n','|') for c in row.cells]
            f.write(f"  Row {ri}: {cells}\n")

print("Done - saved to docx_structure.txt")
